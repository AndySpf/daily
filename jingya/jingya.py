from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches,Cm,Pt
import pandas as pd
import glob
import os
from docxcompose.composer import Composer

filename_list=[]

def read_excel(path, n):
    #df=pd.read_excel(path,header=103)
    #df=pd.read_excel(path,header=0,  skiprows=n+2, nrows= 2)
    df=pd.read_excel(path,header=0,  skiprows=n+2)
    for row in df.values:
        if type(row[2])==float:
            print("未录入，跳过", row[0])
            continue 
        if not row[1]:
            print("异常数据：", row[0])
            break 
        save_img_to_doc(row)
    merge()

def save_img_to_doc(row):
    print(row[1], row[2], row[0])
    WSI_MASK_PATH = os.getcwd() + "/picture/" +str(row[0])#存放图片的文件夹路径
    paths = glob.glob(os.path.join(WSI_MASK_PATH, "*.png"))
    paths.sort()
    simple = True
    if len(paths) > 2:
        simple = False

    """把图片保存到doc文件中的指定位置"""
    tpl_doc = 'simple.docx'
    if not simple:
        tpl_doc = 'double.docx'
    res_doc = "doc/" + str(row[0]) + '.docx'
    # 打开模板文件
    document = Document(tpl_doc)
    table = document.tables[0]
    if simple:
        table.cell(0,1).paragraphs[0].add_run(row[2]).font.size=Pt(16)
        table.cell(0,3).paragraphs[0].add_run(row[1]).font.size=Pt(16)
        table.cell(1,1).paragraphs[0].add_run(str(row[3])).font.size=Pt(16)
        table.cell(1,3).paragraphs[0].add_run(str(row[6]).replace(".0","",1)).font.size=Pt(16)
        table.cell(2,1).paragraphs[0].add_run(row[5]).font.size=Pt(16)
        table.cell(3,1).paragraphs[0].add_run(str(row[4]).replace(".0","",1)).font.size=Pt(16)
        table.cell(3,3).paragraphs[0].add_run(row[7]).font.size=Pt(16)
    else:
        table.cell(0,1).paragraphs[0].add_run(row[2]).font.size=Pt(16)
        table.cell(0,5).paragraphs[0].add_run(row[1]).font.size=Pt(16)
        table.cell(1,1).paragraphs[0].add_run(str(row[3])).font.size=Pt(16)
        table.cell(1,5).paragraphs[0].add_run(str(row[6]).replace(".0","",1)).font.size=Pt(16)
        table.cell(2,1).paragraphs[0].add_run(row[5]).font.size=Pt(16)
        table.cell(3,1).paragraphs[0].add_run(str(row[4]).replace(".0","",1)).font.size=Pt(16)
        table.cell(3,5).paragraphs[0].add_run(row[7]).font.size=Pt(16)
    # 插入图片
    row=int(7)
    col=int(0)
    for path in paths:
        run=table.cell(row,col).paragraphs[0].add_run()
        picture =run.add_picture(path)
        picture.height=Cm(5)
        if not simple:
            picture.width=Cm(5)
        else:
            picture.width=Cm(8)
        col = col + 2
        if simple and col > 3:
            col=0
            row=row+1 
        if not simple and col > 5:
            col=0
            row=row+1
    # 保存结果文件
    document.save(res_doc)
    filename_list.append(res_doc)
    
def merge():
    filename_master = os.getcwd() + "/" + "master.docx" 
    master = Document(filename_master)
    composer = Composer(master)
    for filename in filename_list:
        doc_temp = Document(filename)
#        doc_temp.add_page_break()
        composer.append(doc_temp)
    composer.save(os.getcwd() + "/"+ "new.docx")

def main():
    """主函数"""
    n = input("请输入起始户号：")
    read_excel("source.xlsx", int(n))

if __name__ == '__main__':
    main()

